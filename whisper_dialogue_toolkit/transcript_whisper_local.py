
import argparse
from pathlib import Path
from typing import List, Tuple
from pydub import AudioSegment
from faster_whisper import WhisperModel
from docx import Document

def split_audio(input_path: str, out_dir: Path, chunk_minutes: int, debug: bool=False) -> List[Path]:
    audio = AudioSegment.from_file(input_path)
    total_sec = len(audio) / 1000.0
    if debug:
        print(f"[DEBUG] Loaded audio: {input_path} duration={total_sec:.2f}s")
    chunk_ms = chunk_minutes * 60 * 1000
    chunks = []
    for i, start in enumerate(range(0, len(audio), chunk_ms)):
        end = min(start + chunk_ms, len(audio))
        chunk = audio[start:end]
        dur = len(chunk) / 1000.0
        out_path = out_dir / f"chunk_{i:02d}.wav"
        if dur < 0.2:
            if debug:
                print(f"[DEBUG] Skipping near-empty chunk {i} duration={dur:.3f}s")
            continue
        chunk.export(out_path, format="wav")
        if debug:
            print(f"[DEBUG] Wrote {out_path} duration={dur:.2f}s")
        chunks.append(out_path)
    return chunks

def transcribe_chunk(model, wav_path: Path, language: str="en", vad: bool=False, debug: bool=False):
    segments, info = model.transcribe(
        str(wav_path),
        beam_size=5,
        vad_filter=vad,
        language=None if language == "auto" else language
    )
    out = []
    count = 0
    for seg in segments:
        text = (seg.text or "").strip()
        if text:
            out.append((seg.start, seg.end, text))
            count += 1
    if debug:
        print(f"[DEBUG] Transcribed {wav_path.name}: {count} segments")
    return out

def group_into_turns(segments: List[Tuple[float, float, str]], gap_threshold: float):
    turns = []
    current = []
    last_end = None
    for (s, e, t) in segments:
        if last_end is None:
            current = [(s, e, t)]
        else:
            if s - last_end > gap_threshold:
                if current:
                    turns.append(current)
                current = [(s, e, t)]
            else:
                current.append((s, e, t))
        last_end = e
    if current:
        turns.append(current)
    return turns

def turns_to_dialogue(turns, speaker1: str, speaker2: str):
    lines = []
    for i, turn in enumerate(turns):
        speaker = speaker1 if i % 2 == 0 else speaker2
        text = " ".join(t for (_, _, t) in turn).strip()
        if text:
            lines.append(f"{speaker}: {text}")
    return lines

def write_docx(lines: List[str], out_path: Path):
    doc = Document()
    for line in lines:
        p = doc.add_paragraph()
        if ": " in line:
            speaker, rest = line.split(": ", 1)
            run1 = p.add_run(f"{speaker}: ")
            run1.bold = True
            p.add_run(rest)
        else:
            p.add_run(line)
    doc.save(out_path)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--audio", required=True, help="Path to input audio (mp3/wav/m4a/etc)")
    ap.add_argument("--out_dir", default="out", help="Output directory")
    ap.add_argument("--model", default="medium", help="faster-whisper model: small, medium, large-v3")
    ap.add_argument("--device", default="auto", help="auto|cuda|cpu")
    ap.add_argument("--language", default="en", help="en|auto|<whisper lang code> (default: en)")
    ap.add_argument("--vad", default="false", help="true|false - use VAD filter (default: false)")
    ap.add_argument("--gap_threshold", type=float, default=1.2, help="Silence in seconds to split turns")
    ap.add_argument("--chunk_minutes", type=int, default=10, help="Chunk length in minutes")
    ap.add_argument("--speaker1", default="Anand")
    ap.add_argument("--speaker2", default="Anindya")
    ap.add_argument("--debug", default="true", help="true|false - print debug info")
    args = ap.parse_args()

    debug = str(args.debug).lower() == "true"
    use_vad = str(args.vad).lower() == "true"

    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    chunks_dir = out_dir / "chunks"
    chunks_dir.mkdir(parents=True, exist_ok=True)

    device = args.device if args.device in ("cpu","cuda") else "auto"
    model = WhisperModel(args.model, device=device, compute_type="auto")

    chunks = split_audio(args.audio, chunks_dir, args.chunk_minutes, debug=debug)
    if not chunks:
        print("[ERROR] No non-empty chunks were produced. Ensure ffmpeg is installed and the audio file is valid.")
        return

    all_turns = []
    for wav in chunks:
        segments = transcribe_chunk(model, wav, language=args.language, vad=use_vad, debug=debug)
        if not segments:
            if debug:
                print(f"[DEBUG] No speech recognized in {wav.name}, skipping.")
            continue
        raw_txt = "\n".join([f"[{s:.2f}-{e:.2f}] {t}" for (s,e,t) in segments])
        (chunks_dir / f"{wav.stem}.txt").write_text(raw_txt, encoding="utf-8")
        turns = group_into_turns(segments, args.gap_threshold)
        all_turns.extend(turns)

    if not all_turns:
        print("[ERROR] No turns were detected. Try --language auto, set --vad false, or reduce --gap_threshold.")
        return

    lines = turns_to_dialogue(all_turns, args.speaker1, args.speaker2)
    (out_dir / "dialogue.txt").write_text("\n".join(lines), encoding="utf-8")
    write_docx(lines, out_dir / "dialogue.docx")

    print(f"Done. Outputs in: {out_dir}")
    print(f"- {out_dir / 'dialogue.txt'}")
    print(f"- {out_dir / 'dialogue.docx'}")
    print(f"- Chunk raw segments in: {chunks_dir}")

if __name__ == "__main__":
    main()
